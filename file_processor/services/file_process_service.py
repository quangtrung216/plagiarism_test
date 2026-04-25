from pathlib import Path # Ko cần quan tâm file này
from typing import Optional

try:
    import fitz
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

from nlp.text_preprocess import (
    extract_plagiarism_zone,
    normalize_full_text,
    preprocess_pdf_text,
)


class FileProcessService:
    """Service for extracting text from uploaded files and running preprocessing."""

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> dict:
        if not HAS_PYMUPDF:
            return {
                "success": False,
                "error": "PyMuPDF not installed. Run: pip install PyMuPDF",
                "file_type": "pdf",
            }

        try:
            text_chunks = []
            with fitz.open(file_path) as pdf:
                page_count = pdf.page_count
                for page in pdf:
                    text_chunks.append(page.get_text("text") or "")

            return {
                "success": True,
                "text": "\n".join(text_chunks),
                "pages": page_count,
                "file_type": "pdf",
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "file_type": "pdf",
            }

    @staticmethod
    def extract_text_from_word(file_path: str) -> dict:
        if not HAS_PYTHON_DOCX:
            return {
                "success": False,
                "error": "python-docx not installed. Run: pip install python-docx",
                "file_type": "docx",
            }

        try:
            doc = Document(file_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text:
                    text_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_parts.append(cell.text)

            return {
                "success": True,
                "text": "\n".join(text_parts),
                "paragraphs": len(doc.paragraphs),
                "file_type": "docx",
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "file_type": "docx",
            }

    @staticmethod
    def extract_text_from_file(file_path: str, file_type: Optional[str] = None) -> dict:
        if file_type is None:
            file_ext = Path(file_path).suffix.lower()
            if file_ext == ".pdf":
                file_type = "pdf"
            elif file_ext == ".docx":
                file_type = "docx"
            else:
                return {"success": False, "error": f"Unsupported file type: {file_ext}"}

        normalized_file_type = file_type.lower()

        if normalized_file_type == "pdf":
            return FileProcessService.extract_text_from_pdf(file_path)

        if normalized_file_type in {"word", "docx"}:
            return FileProcessService.extract_text_from_word(file_path)

        return {"success": False, "error": f"Unsupported file type: {file_type}"}

    @staticmethod
    def process_file_and_extract_sentences(
        file_path: str,
        language: str = "vi",
        min_words: int = 8,
        max_words: int = 200,
        use_zone: bool = True,
        file_type: Optional[str] = None,
    ) -> dict:
        extraction_result = FileProcessService.extract_text_from_file(file_path, file_type)
        if not extraction_result["success"]:
            return extraction_result

        extracted_text = extraction_result.get("text", "")

        try:
            normalized_text = normalize_full_text(extracted_text)
            plagiarism_zone = (
                extract_plagiarism_zone(normalized_text) if use_zone else normalized_text
            )

            sentences = preprocess_pdf_text(
                raw_text=plagiarism_zone,
                language=language,
                min_words=min_words,
                max_words=max_words,
                use_zone=False,
            )

            return {
                "success": True,
                "file_type": extraction_result.get("file_type"),
                "total_sentences": len(sentences),
                "sentences": sentences,
                "metadata": {
                    "language": language,
                    "min_words": min_words,
                    "max_words": max_words,
                    "use_zone": use_zone,
                    "original_text_length": len(extracted_text),
                    "normalized_text_length": len(normalized_text),
                    "plagiarism_zone_length": len(plagiarism_zone),
                    "pages_or_paragraphs": extraction_result.get("pages")
                    or extraction_result.get("paragraphs", 0),
                },
            }
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "file_type": extraction_result.get("file_type"),
            }
